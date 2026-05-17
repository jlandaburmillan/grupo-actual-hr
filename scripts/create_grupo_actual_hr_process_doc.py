from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/jorgelandaburmillan/Documents/Codex/grupo-actual-hr-final/Grupo_Actual_HR_Documentacion_Proceso.docx"

NAVY = "0A3761"
BLUE = "0A6ED1"
CYAN = "00B8D9"
GREY = "F5F7FA"
GRAPHITE = "1D2D3E"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2EC", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, GREY)
    set_cell_border(cell, "B8DDEB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    add_run(p2, body, color=GRAPHITE)
    doc.add_paragraph()


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2.15, 4.35])
    hdr = table.rows[0].cells
    hdr[0].text = "Elemento"
    hdr[1].text = "Definición"
    for cell in hdr:
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE)
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        set_cell_shading(cells[0], GREY)
        for cell in cells:
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor.from_string(GRAPHITE)
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()
    return table


def add_three_col_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [1.75, 2.45, 2.30])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE)
                run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_border(cells[idx])
            if idx == 0:
                set_cell_shading(cells[idx], GREY)
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor.from_string(GRAPHITE)
                    if idx == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(GRAPHITE)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in [
    ("Title", 24, NAVY, 0, 10),
    ("Subtitle", 12, GRAPHITE, 0, 14),
    ("Heading 1", 16, NAVY, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = style_name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.10

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(header, "Grupo Actual HR · Documentación de proceso", color=NAVY, size=8.5)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(footer, "Blueprint funcional para Capital Humano", color=BLUE, size=8.5)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(title, "Grupo Actual HR", bold=True, color=NAVY, size=26)
subtitle = doc.add_paragraph(style="Subtitle")
add_run(subtitle, "Documentación del proceso completo de diseño, rebranding, construcción y validación de la landing page corporativa.", color=GRAPHITE)

add_callout(
    doc,
    "Resultado validado",
    "La versión final quedó operativa en http://127.0.0.1:3022/ y el build de producción fue validado correctamente con npm run build."
)

add_heading(doc, "1. Resumen ejecutivo", 1)
doc.add_paragraph(
    "El proyecto comenzó como una landing corporativa para una marca previa y evolucionó hacia una identidad completamente renovada: Grupo Actual HR. "
    "La solución final posiciona a la firma como consultora enterprise especializada en procesos de Capital Humano, Discovery funcional, arquitectura de procesos, SAP HCM and SAP SuccessFactors, trazabilidad, cumplimiento y adopción operacional."
)
doc.add_paragraph(
    "El foco estratégico dejó de ser una narrativa genérica de tecnología HR Cloud y pasó a una propuesta consultiva basada en procesos: antes de implementar tecnología, Grupo Actual HR entiende el negocio, diseña el proceso y luego habilita la solución."
)

add_heading(doc, "2. Punto de partida y cambio de dirección", 1)
add_bullet(doc, "La primera base se trabajó como landing B2B para una marca distinta, con servicios SAP HCM, SAP SuccessFactors y WorkForce Software.")
add_bullet(doc, "Luego se solicitó un rebranding completo hacia Grupo Actual HR, eliminando referencias anteriores y cambiando el eje visual a Blueprint funcional para Capital Humano.")
add_bullet(doc, "Se definió una estética SAP-like consultiva, más cercana a una firma de procesos enterprise que a una empresa de software genérica.")
add_bullet(doc, "Se construyó una versión limpia y validada en la carpeta final del proyecto.")

add_heading(doc, "3. Brand kit vigente", 1)
add_kv_table(doc, [
    ("Nombre oficial", "Grupo Actual HR"),
    ("Logo textual", "GRUPO arriba y Actual HR abajo, sin isotipo dominante."),
    ("Eje conceptual", "Blueprint funcional para Capital Humano."),
    ("Territorio visual", "Ruta Blueprint."),
    ("Promesa", "Diseñamos procesos de Capital Humano para implementar tecnología con sentido de negocio."),
    ("Frase estratégica", "Antes de implementar, entendemos. Antes de configurar, diseñamos. Antes del sistema, está el proceso."),
    ("Personalidad", "Consultiva, técnica, enterprise, humana y clara."),
])

add_heading(doc, "4. Paleta aprobada", 1)
add_kv_table(doc, [
    ("Deep SAP Navy", "#0A3761 - logo, fondos institucionales, títulos y autoridad enterprise."),
    ("SAP-like Blue", "#0A6ED1 - acentos principales, navegación, botones y links."),
    ("Electric Cyan", "#00B8D9 - nodos, líneas Blueprint, highlights y trazabilidad."),
    ("Cloud White", "#FFFFFF - fondos principales y claridad visual."),
    ("Fiori Grey", "#F5F7FA - fondos secundarios, tarjetas y bloques."),
    ("Graphite Blue", "#1D2D3E - textos, profundidad y contraste ejecutivo."),
])

add_heading(doc, "5. Decisiones editoriales", 1)
add_bullet(doc, "Todo el contenido visible del sitio se mantuvo en español.")
add_bullet(doc, "Cuando se conectan ambas plataformas, se usa la forma SAP HCM and SAP SuccessFactors.")
add_bullet(doc, "Se reemplazaron claims en inglés visibles por expresiones comerciales en español, como Blueprint funcional para Capital Humano y Ruta Blueprint.")
add_bullet(doc, "Se eliminó la dependencia conceptual de la marca anterior y se mantuvo una narrativa consultiva de procesos.")

add_heading(doc, "6. Arquitectura de la landing", 1)
add_three_col_table(doc, ["Sección", "Objetivo", "Resultado"], [
    ("Header", "Navegación sticky con logo textual.", "Contraste corregido y CTA visible."),
    ("Hero", "Presentar propuesta central y visual Blueprint.", "Incluye nodos, capas, BPML y portlets estratégicos."),
    ("Esencia de marca", "Explicar convicción de procesos antes de tecnología.", "Dos cards: Procesos y Consultoría."),
    ("Blueprint", "Mostrar método funcional de trabajo.", "Cuatro pasos: Entender, Diseñar, Implementar, Adoptar."),
    ("Servicios", "Ordenar líneas comerciales.", "Cards con capacidades y microcopy consultivo."),
    ("SAP HCM", "Detallar continuidad, control y evolución.", "Lista funcional actualizada por solicitud del usuario."),
    ("SAP SuccessFactors", "Explicar transformación cloud con sentido de negocio.", "Módulos y mensaje destacado."),
    ("Metodología", "Presentar ciclo de trabajo.", "Pasos en español: Descubrir, Mapear, Diseñar, Habilitar, Adoptar, Mejorar."),
    ("CTA final", "Convertir interés en contacto.", "Formulario visual y CTA de diagnóstico."),
])

add_heading(doc, "7. Cambios visuales clave", 1)
add_bullet(doc, "Se removió el isotipo del logo para mantener únicamente GRUPO / Actual HR.")
add_bullet(doc, "Se corrigió el contraste del menú principal para asegurar legibilidad sobre fondo institucional.")
add_bullet(doc, "Se reforzó el hero con fondo Deep SAP Navy, patrón Blueprint y panel de arquitectura funcional HR.")
add_bullet(doc, "Se incorporaron portlets específicos: colaborador, cumplimiento, automatización, sistemas y temas críticos de implementación.")
add_bullet(doc, "Se agregó la bajada BPML (Maestro de Procesos).")
add_bullet(doc, "Se corrigió Tailwind agregando el alias deep-navy, necesario para cargar correctamente el azul institucional en producción.")

add_heading(doc, "8. Stack técnico", 1)
add_kv_table(doc, [
    ("Framework", "Next.js 14.2.35 con App Router."),
    ("Lenguaje", "TypeScript."),
    ("UI", "React con componentes reutilizables por sección."),
    ("Estilos", "Tailwind CSS con paleta personalizada."),
    ("Íconos", "lucide-react."),
    ("Deploy objetivo", "Vercel."),
    ("Carpeta final", "/Users/jorgelandaburmillan/Documents/Codex/grupo-actual-hr-final"),
])

add_heading(doc, "9. Archivos principales", 1)
add_bullet(doc, "app/layout.tsx - metadata SEO y estructura raíz.")
add_bullet(doc, "app/page.tsx - composición de secciones de la landing.")
add_bullet(doc, "app/globals.css - sistema visual base, spacing, patrones y cards.")
add_bullet(doc, "components/Header.tsx - logo, navegación y menú mobile.")
add_bullet(doc, "components/Hero.tsx - propuesta principal y visual Blueprint.")
add_bullet(doc, "components/Services.tsx - líneas de servicio.")
add_bullet(doc, "components/SapHcmSection.tsx - sección SAP HCM actualizada.")
add_bullet(doc, "components/Methodology.tsx - metodología en español.")
add_bullet(doc, "tailwind.config.ts - paleta oficial, alias deep-navy y sombras.")
add_bullet(doc, "GRUPO_ACTUAL_HR_HANDOFF_PROMPT.md - prompt maestro actualizado.")
add_bullet(doc, "PROMPT_APP_WEB_RESPONSIVA_GRUPO_ACTUAL_HR.md - prompt para evolucionar a app web.")

add_heading(doc, "10. Validaciones realizadas", 1)
add_bullet(doc, "Se instaló una base limpia de dependencias.")
add_bullet(doc, "Se corrigió la dependencia faltante lucide-react.")
add_bullet(doc, "Se validó TypeScript sin errores.")
add_bullet(doc, "Se ejecutó npm run build correctamente.")
add_bullet(doc, "Se levantó servidor local de producción.")
add_bullet(doc, "Se verificó respuesta HTTP 200.")
add_bullet(doc, "Se corrigió el problema visual de carga parcial agregando el alias deep-navy en Tailwind.")
add_bullet(doc, "Se confirmó visualmente que la página queda desplegada correctamente en http://127.0.0.1:3022/.")

add_heading(doc, "11. Entregables generados", 1)
add_kv_table(doc, [
    ("Sitio final", "/Users/jorgelandaburmillan/Documents/Codex/grupo-actual-hr-final"),
    ("Zip de entrega", "/Users/jorgelandaburmillan/Documents/Codex/grupo-actual-hr-final.zip"),
    ("Prompt maestro", "GRUPO_ACTUAL_HR_HANDOFF_PROMPT.md"),
    ("Prompt app web", "PROMPT_APP_WEB_RESPONSIVA_GRUPO_ACTUAL_HR.md"),
    ("Documento Word", "Grupo_Actual_HR_Documentacion_Proceso.docx"),
])

add_heading(doc, "12. Próximos pasos recomendados", 1)
for item in [
    "Subir el proyecto a un repositorio Git.",
    "Conectar el repositorio a Vercel.",
    "Configurar dominio final.",
    "Implementar formulario real de contacto.",
    "Separar contenido comercial en un archivo editable de contenido.",
    "Agregar sitemap, robots, Open Graph y analítica.",
    "Crear una ruta para brochure o material descargable.",
    "Preparar política de privacidad si el formulario captura datos reales.",
]:
    add_number(doc, item)

add_callout(
    doc,
    "Conclusión",
    "Grupo Actual HR quedó documentado, validado y preparado como base de producción. La siguiente etapa natural es convertir la landing en una app web con formulario funcional, rutas adicionales y despliegue en Vercel."
)

doc.save(OUT)
print(OUT)
